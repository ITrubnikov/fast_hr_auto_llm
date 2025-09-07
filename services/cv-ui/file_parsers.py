"""
Модуль для парсинга различных форматов файлов резюме
Поддерживает: PDF, DOC, DOCX, TXT, RTF, HTML, MD
Архитектура согласно SUPPORTED_FORMATS.md
"""

import os
import io
import re
from typing import Optional, List, Dict, Any
from pathlib import Path
import logging
from abc import ABC, abstractmethod

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class BaseParser(ABC):
    """Базовый абстрактный класс для всех парсеров"""
    
    @abstractmethod
    def parse(self, file_path: str) -> str:
        """Парсит файл и возвращает извлеченный текст"""
        pass
    
    @abstractmethod
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные файла"""
        pass
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Очищает извлеченный текст от лишних символов"""
        if not text:
            return ""
        
        # Удаляем лишние пробелы и переносы строк
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line and not line.isspace():
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)

class PDFParser(BaseParser):
    """Парсер для PDF файлов через OCR"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """Парсит PDF файл через OCR"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            return BaseParser.clean_text(text)
            
        except ImportError:
            logger.error("PyMuPDF не установлен. Установите: pip install PyMuPDF")
            raise
        except Exception as e:
            logger.error(f"Ошибка при парсинге PDF {file_path}: {e}")
            raise
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные PDF"""
        try:
            import fitz
            
            doc = fitz.open(file_path)
            metadata = doc.metadata
            page_count = len(doc)
            doc.close()
            
            return {
                "title": metadata.get("title", ""),
                "author": metadata.get("author", ""),
                "subject": metadata.get("subject", ""),
                "creator": metadata.get("creator", ""),
                "producer": metadata.get("producer", ""),
                "creation_date": metadata.get("creationDate", ""),
                "modification_date": metadata.get("modDate", ""),
                "pages": page_count
            }
        except Exception as e:
            logger.error(f"Ошибка при извлечении метаданных PDF {file_path}: {e}")
            return {}

class WordParser(BaseParser):
    """Парсер для DOC/DOCX файлов"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """Парсит Word документ"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.docx':
            return WordParser._parse_docx(file_path)
        elif ext == '.doc':
            return WordParser._parse_doc(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат Word: {ext}")
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Парсит DOCX файл"""
        try:
            # Проверяем, что это действительно DOCX файл
            import zipfile
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Если это ZIP файл, то это DOCX
                    pass
            except zipfile.BadZipFile:
                # Если не ZIP, то это может быть PDF с неправильным расширением
                logger.warning(f"Файл {file_path} не является DOCX, пытаемся парсить как PDF")
                return PDFParser.parse(file_path)
            
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return BaseParser.clean_text(text)
            
        except ImportError:
            logger.error("python-docx не установлен. Установите: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Ошибка при парсинге DOCX {file_path}: {e}")
            # Fallback на PDF парсер
            try:
                return PDFParser.parse(file_path)
            except Exception as pdf_error:
                logger.error(f"Ошибка при fallback парсинге как PDF {file_path}: {pdf_error}")
                raise
    
    @staticmethod
    def _parse_doc(file_path: str) -> str:
        """Парсит DOC файл (старый формат)"""
        try:
            # Проверяем, не является ли файл PDF (по первым байтам)
            with open(file_path, 'rb') as f:
                first_bytes = f.read(10)
                if first_bytes.startswith(b'%PDF'):
                    logger.warning(f"Файл {file_path} имеет расширение .doc, но является PDF. Парсим как PDF.")
                    return PDFParser.parse(file_path)
            
            # Для старых DOC файлов используем python-docx2txt
            import docx2txt
            
            text = docx2txt.process(file_path)
            return BaseParser.clean_text(text)
            
        except ImportError:
            logger.error("docx2txt не установлен. Установите: pip install docx2txt")
            raise
        except Exception as e:
            logger.error(f"Ошибка при парсинге DOC {file_path}: {e}")
            # Fallback на простое чтение как текстового файла
            return WordParser._fallback_doc_parse(file_path)
    
    @staticmethod
    def _fallback_doc_parse(file_path: str) -> str:
        """Fallback парсинг DOC файла как текстового"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Простое удаление бинарных данных
            import re
            text = re.sub(r'[^\x20-\x7E\n\r\t]', '', content)
            
            return BaseParser.clean_text(text)
        except Exception as e:
            logger.error(f"Ошибка при fallback парсинге DOC {file_path}: {e}")
            raise
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные Word документа"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            core_props = doc.core_properties
            
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or ""
            }
        except Exception as e:
            logger.error(f"Ошибка при извлечении метаданных Word {file_path}: {e}")
            return {}

class TextParser(BaseParser):
    """Парсер для текстовых файлов (TXT)"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """Парсит текстовый файл с поддержкой различных кодировок"""
        encodings = ['utf-8', 'cp1251', 'windows-1251', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    return BaseParser.clean_text(text)
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Ошибка при чтении файла {file_path} с кодировкой {encoding}: {e}")
                continue
        
        raise ValueError(f"Не удалось прочитать файл {file_path} ни с одной из поддерживаемых кодировок")
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные текстового файла"""
        try:
            stat = os.stat(file_path)
            return {
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "encoding": "utf-8"  # Будет определено при чтении
            }
        except Exception as e:
            logger.error(f"Ошибка при извлечении метаданных TXT {file_path}: {e}")
            return {}

class RTFParser(BaseParser):
    """Парсер для RTF файлов"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """Парсит RTF файл"""
        try:
            import striprtf
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_content = file.read()
            
            text = striprtf.rtf_to_text(rtf_content)
            return BaseParser.clean_text(text)
            
        except ImportError:
            logger.error("striprtf не установлен. Установите: pip install striprtf")
            raise
        except Exception as e:
            logger.error(f"Ошибка при парсинге RTF {file_path}: {e}")
            # Fallback на простой парсинг
            return RTFParser._fallback_parse(file_path)
    
    @staticmethod
    def _fallback_parse(file_path: str) -> str:
        """Fallback парсинг RTF без библиотеки striprtf"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Простое удаление RTF тегов
            import re
            text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\\[{}]', '', text)
            
            return BaseParser.clean_text(text)
        except Exception as e:
            logger.error(f"Ошибка при fallback парсинге RTF {file_path}: {e}")
            raise
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные RTF файла"""
        try:
            stat = os.stat(file_path)
            return {
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "format": "RTF"
            }
        except Exception as e:
            logger.error(f"Ошибка при извлечении метаданных RTF {file_path}: {e}")
            return {}

class HTMLParser(BaseParser):
    """Парсер для HTML файлов"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """Парсит HTML файл"""
        try:
            from bs4 import BeautifulSoup
            import html2text
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
            
            # Используем BeautifulSoup для очистки
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Удаляем скрипты и стили
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Извлекаем текст
            text = soup.get_text()
            
            # Дополнительная очистка через html2text
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.ignore_images = True
            h.ignore_emphasis = True
            
            clean_text = h.handle(html_content)
            
            # Используем результат BeautifulSoup как основной, html2text как fallback
            if text.strip():
                return BaseParser.clean_text(text)
            else:
                return BaseParser.clean_text(clean_text)
                
        except ImportError:
            logger.error("BeautifulSoup4 или html2text не установлены. Установите: pip install beautifulsoup4 html2text")
            raise
        except Exception as e:
            logger.error(f"Ошибка при парсинге HTML {file_path}: {e}")
            # Fallback на простой парсинг
            return HTMLParser._fallback_parse(file_path)
    
    @staticmethod
    def _fallback_parse(file_path: str) -> str:
        """Fallback парсинг HTML без библиотек"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Простое удаление HTML тегов
            import re
            text = re.sub(r'<[^>]+>', '', content)
            text = re.sub(r'&[a-zA-Z0-9#]+;', ' ', text)  # HTML entities
            
            return BaseParser.clean_text(text)
        except Exception as e:
            logger.error(f"Ошибка при fallback парсинге HTML {file_path}: {e}")
            raise
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные HTML файла"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            return {
                "title": soup.title.string if soup.title else "",
                "meta_description": "",
                "meta_keywords": "",
                "meta_author": "",
                "charset": "utf-8"
            }
        except Exception as e:
            logger.error(f"Ошибка при извлечении метаданных HTML {file_path}: {e}")
            return {}

class MarkdownParser(BaseParser):
    """Парсер для Markdown файлов"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """Парсит Markdown файл"""
        try:
            import markdown
            
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Конвертируем в HTML, затем в текст
            html = markdown.markdown(md_content)
            
            # Удаляем HTML теги
            import re
            text = re.sub(r'<[^>]+>', '', html)
            
            return BaseParser.clean_text(text)
            
        except ImportError:
            logger.error("markdown не установлен. Установите: pip install markdown")
            raise
        except Exception as e:
            logger.error(f"Ошибка при парсинге Markdown {file_path}: {e}")
            # Fallback на простое чтение
            return MarkdownParser._fallback_parse(file_path)
    
    @staticmethod
    def _fallback_parse(file_path: str) -> str:
        """Fallback парсинг Markdown без библиотеки"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Простое удаление Markdown разметки
            import re
            text = re.sub(r'#{1,6}\s+', '', content)  # Заголовки
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Жирный текст
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Курсив
            text = re.sub(r'`(.*?)`', r'\1', text)  # Код
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Ссылки
            text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)  # Списки
            text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)  # Нумерованные списки
            
            return BaseParser.clean_text(text)
        except Exception as e:
            logger.error(f"Ошибка при fallback парсинге Markdown {file_path}: {e}")
            raise
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные Markdown файла"""
        try:
            stat = os.stat(file_path)
            return {
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "format": "Markdown"
            }
        except Exception as e:
            logger.error(f"Ошибка при извлечении метаданных Markdown {file_path}: {e}")
            return {}

class FileParser:
    """Главный класс для парсинга файлов резюме"""
    
    # Регистр парсеров
    _parsers = {
        '.pdf': PDFParser,
        '.doc': WordParser,
        '.docx': WordParser,
        '.txt': TextParser,
        '.rtf': RTFParser,
        '.html': HTMLParser,
        '.htm': HTMLParser,
        '.md': MarkdownParser
    }
    
    @staticmethod
    def get_supported_extensions() -> List[str]:
        """Возвращает список поддерживаемых расширений файлов"""
        return list(FileParser._parsers.keys())
    
    @staticmethod
    def is_supported_file(file_path: str) -> bool:
        """Проверяет, поддерживается ли формат файла"""
        ext = Path(file_path).suffix.lower()
        return ext in FileParser._parsers
    
    @staticmethod
    def parse_file(file_path: str) -> str:
        """
        Парсит файл и возвращает извлеченный текст
        Автоматически определяет формат по расширению
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        
        if ext not in FileParser._parsers:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")
        
        parser_class = FileParser._parsers[ext]
        return parser_class.parse(file_path)
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает метаданные файла"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        
        if ext not in FileParser._parsers:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")
        
        parser_class = FileParser._parsers[ext]
        return parser_class.get_metadata(file_path)
    
    @staticmethod
    def parse_file_with_metadata(file_path: str) -> Dict[str, Any]:
        """Парсит файл и возвращает текст с метаданными"""
        text = FileParser.parse_file(file_path)
        metadata = FileParser.get_metadata(file_path)
        
        return {
            "text": text,
            "metadata": metadata,
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "file_extension": Path(file_path).suffix.lower(),
            "file_size": len(text)
        }

# Функция для обратной совместимости
def extract_text_from_file(file_path: str) -> str:
    """Извлекает текст из файла (обратная совместимость)"""
    return FileParser.parse_file(file_path)

# Функции извлечения данных из текста резюме
def extract_phone_from_text(text: str) -> str:
    """
    Извлекает номер телефона из текста резюме
    Поддерживает различные форматы:
    - +7 903 333-33-33
    - +7 (927) 1212210
    - 8 903 333-33-33
    - 903 333-33-33
    - +7(903)333-33-33
    - 8-903-333-33-33
    """
    if not text:
        return "телефон не найден"
    
    # Паттерны для различных форматов российских номеров (в порядке приоритета)
    phone_patterns = [
        # +7 (927) 1212210 - приоритетный формат
        r'\+7\s*\((\d{3})\)\s*(\d{7})',
        # +7 903 333-33-33, +7 (903) 333-33-33
        r'\+7\s*\(?(\d{3})\)?\s*(\d{3})[\s\-]?(\d{2})[\s\-]?(\d{2})',
        # 8 (927) 1212210
        r'8\s*\((\d{3})\)\s*(\d{7})',
        # 8 903 333-33-33, 8 (903) 333-33-33
        r'8\s*\(?(\d{3})\)?\s*(\d{3})[\s\-]?(\d{2})[\s\-]?(\d{2})',
        # (927) 1212210
        r'\((\d{3})\)\s*(\d{7})',
        # 903 333-33-33, (903) 333-33-33
        r'\(?(\d{3})\)?\s*(\d{3})[\s\-]?(\d{2})[\s\-]?(\d{2})',
        # +7 903 333 33 33
        r'\+7\s*(\d{3})\s*(\d{3})\s*(\d{2})\s*(\d{2})',
        # 8 903 333 33 33
        r'8\s*(\d{3})\s*(\d{3})\s*(\d{2})\s*(\d{2})',
        # 903 333 33 33
        r'(\d{3})\s*(\d{3})\s*(\d{2})\s*(\d{2})',
        # +7(903)333-33-33
        r'\+7\((\d{3})\)(\d{3})-(\d{2})-(\d{2})',
        # 8(903)333-33-33
        r'8\((\d{3})\)(\d{3})-(\d{2})-(\d{2})',
        # (903)333-33-33
        r'\((\d{3})\)(\d{3})-(\d{2})-(\d{2})',
        # 8-903-333-33-33
        r'8-(\d{3})-(\d{3})-(\d{2})-(\d{2})',
        # 903-333-33-33
        r'(\d{3})-(\d{3})-(\d{2})-(\d{2})',
        # Простые 10-11 цифр подряд (только если начинается с 7 или 8)
        r'\b(7|8)\d{10}\b',
        # 7 цифр подряд (без кода страны) - только если в контексте телефона
        r'(?:телефон|phone|тел\.?)\s*:?\s*(\d{3})\s*(\d{2})\s*(\d{2})'
    ]
    
    # Ищем все возможные номера
    found_phones = []
    
    for pattern in phone_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                # Собираем номер из групп
                phone_parts = [part for part in match if part]
                if len(phone_parts) >= 2:  # Минимум 2 группы цифр
                    phone = ''.join(phone_parts)
                    if len(phone) >= 7:  # Минимум 7 цифр
                        found_phones.append(phone)
            else:
                if len(match) >= 7:
                    found_phones.append(match)
    
    # Убираем дубликаты и выбираем по приоритету
    unique_phones = list(set(found_phones))
    if unique_phones:
        # Приоритет: сначала номера с +7, потом с 8, потом остальные
        def phone_priority(phone):
            if phone.startswith('+7') or phone.startswith('7'):
                return 0
            elif phone.startswith('8'):
                return 1
            else:
                return 2
        
        unique_phones.sort(key=lambda x: (phone_priority(x), -len(x)))
        return unique_phones[0]
    
    return "телефон не найден"

def extract_email_from_text(text: str) -> str:
    """Извлекает email из текста резюме"""
    if not text:
        return "email не найден"
    
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else "email не найден"

def extract_name_from_text(text: str) -> str:
    """Извлекает имя из текста резюме"""
    if not text:
        return "Имя не найдено"
    
    lines = text.split('\n')
    for line in lines[:10]:  # Ищем в первых 10 строках
        line = line.strip()
        if len(line) > 3 and len(line) < 50:
            # Проверяем, что это похоже на имя (содержит буквы и пробелы)
            if re.match(r'^[А-Яа-яЁё\s]+$', line) and len(line.split()) >= 2:
                # Исключаем служебные слова
                if not any(word in line.lower() for word in ['резюме', 'зарплата', 'телефон', 'email', 'возраст', 'дата', 'опыт', 'образование', 'навыки']):
                    return line
    
    return "Имя не найдено"

def extract_data_from_cv(text: str) -> Dict[str, str]:
    """Извлекает основные данные из текста резюме"""
    return {
        "name": extract_name_from_text(text),
        "email": extract_email_from_text(text),
        "phone": extract_phone_from_text(text)
    }